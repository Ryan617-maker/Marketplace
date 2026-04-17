import json  # json
from pathlib import Path

import pytest
from docx import Document  # docx
from datetime import datetime  # data


def rodar_testes():
    class ColetorResultados:
        def __init__(self):
            self.tests = []

        def pytest_runtest_logreport(self, report):
            # Considera apenas o resultado final de cada caso de teste, usando dicionário para filtrar.
            if report.when == "call":
                self.tests.append({"nodeid": report.nodeid, "outcome": report.outcome})

    coletor = ColetorResultados()
    codigo_saida = pytest.main(["."], plugins=[coletor])

    summary = {
        "total": len(coletor.tests),
        "passed": sum(1 for t in coletor.tests if t["outcome"] == "passed"),
        "failed": sum(1 for t in coletor.tests if t["outcome"] == "failed"),
    }

    with open("report.json", "w", encoding="utf-8") as f:
        json.dump(
            {"summary": summary, "tests": coletor.tests},
            f,
            ensure_ascii=False,
            indent=2,
        )

    return codigo_saida


def carregar_relatorio():
    caminho_relatorio = Path("report.json")
    if not caminho_relatorio.exists():
        raise FileNotFoundError(
            "Arquivo report.json não encontrado. Verifique se os testes foram executados com sucesso."
        )

    with caminho_relatorio.open(encoding="utf-8") as f:
        return json.load(f)


def extrair_dados(dados):
    total = dados["summary"]["total"]
    passed = dados["summary"]["passed"]
    failed = dados["summary"]["failed"]

    testes = []

    for teste in dados["tests"]:
        testes.append({"nome": teste["nodeid"], "status": teste["outcome"]})

    return total, passed, failed, testes


def criar_doc(total, passed, failed, testes):
    doc = Document()

    doc.add_heading("RELATÓRIO DE TESTES DE SOFTWARE", 0)

    doc.add_heading("1. Informações da Execução", level=1)
    doc.add_paragraph(f'Data: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    doc.add_paragraph("Ambiente: Local")
    doc.add_paragraph("Ferramenta de Testes: Pytest")

    doc.add_heading("2. Resultados Gerais", level=1)
    doc.add_paragraph(f"Total de testes: {total}")
    doc.add_paragraph(f"Testes aprovados: {passed}")
    doc.add_paragraph(f"Testes com falha: {failed}")

    status_geral = "APROVADO" if failed == 0 else "REPROVADO"
    doc.add_paragraph(f"Status da execução: {status_geral}")

    doc.add_heading("3. Detalhamento dos Testes", level=1)

    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Table Grid"

    header = tabela.rows[0].cells
    header[0].text = "Caso de Teste"
    header[1].text = "Resultado"

    for t in testes:
        row = tabela.add_row().cells
        row[0].text = t["nome"]
        row[1].text = "Aprovado" if t["status"] == "passed" else "Falhou"

    doc.add_heading("4. Considerações Finais", level=1)

    if failed == 0:
        doc.add_paragraph("Todos os testes foram executados com sucesso.")
    else:
        doc.add_paragraph(
            "Foram identificadas falhas que devem ser analisadas e corrigidas."
        )

    doc.save("relatorio_testes.docx")


if __name__ == "__main__":
    rodar_testes()
    dados = carregar_relatorio()
    total, passed, failed, testes = extrair_dados(dados)
    criar_doc(total, passed, failed, testes)
